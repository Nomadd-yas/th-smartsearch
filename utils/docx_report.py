"""Генерация DOCX-документа обоснования НМЦК по шаблону МР №567.

Структура документа строго соответствует шаблону «обоснование нмцк.docx»:
  • Реестровый номер закупки (генерируется автоматически)
  • Правовое основание (ФЗ-44, МР №567)
  • Для каждой позиции: таблица контрактов + расчёт НМЦК
  • Итог: коэффициент вариации, валюта, подпись

Публичный API:
  build_nmck_docx(items, registry_number) → bytes
"""

from __future__ import annotations

import io
import statistics as _stats
from datetime import date as _date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ── Константы форматирования ──────────────────────────────────────

_FONT   = "Times New Roman"
_SIZE   = Pt(14)       # ГОСТ для официальных документов — 14pt
_SIZE_T = Pt(12)       # таблицы


# ── Низкоуровневые хелперы ────────────────────────────────────────

def _set_font(run, size: Pt = _SIZE, bold: bool = False) -> None:
    run.font.name = _FONT
    run.font.size = size
    run.font.bold = bold
    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs")):
        rFonts.set(attr, _FONT)


def _spacing(para, before: int = 0, after: int = 0, line: int | None = None) -> None:
    pPr  = para._p.get_or_add_pPr()
    spec = OxmlElement("w:spacing")
    spec.set(qn("w:before"), str(before))
    spec.set(qn("w:after"),  str(after))
    if line:
        spec.set(qn("w:line"),     str(line))
        spec.set(qn("w:lineRule"), "auto")
    pPr.append(spec)


def _add_para(doc: Document, text: str = "", bold: bool = False,
              align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
              size: Pt = _SIZE, before: int = 0, after: int = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    _spacing(p, before=before, after=after)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold)


def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _cell(cell, text: str, bold: bool = False,
          align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
          size: Pt = _SIZE_T) -> None:
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold)


# ── Числовой формат ───────────────────────────────────────────────

def _fmt(value: float) -> str:
    """1 234,56 (без знака валюты — для таблиц)"""
    return f"{value:,.2f}".replace(",", "\u00a0").replace(".", ",")


def _ru(value: float) -> str:
    """1 234,56 руб."""
    return f"{_fmt(value)} руб."


def _pct(data: list[float], p: float) -> float:
    if not data:
        return 0.0
    n   = len(data)
    idx = p / 100 * (n - 1)
    lo, hi = int(idx), min(int(idx) + 1, n - 1)
    return data[lo] + (data[hi] - data[lo]) * (idx - lo)


# ── Блок одной позиции ────────────────────────────────────────────

def _item_block(doc: Document, item: dict, show_name: bool = True) -> None:
    """Таблица контрактов + расчёт НМЦК для одной позиции."""

    nmck_data   = item.get("nmck_data") or {}
    all_ann     = nmck_data.get("contracts", [])
    used        = [
        c["contract"] for c in all_ann
        if c.get("status") in ("used", "force_included")
    ]

    unit_price = item.get("unit_price", 0.0)
    quantity   = item.get("quantity") or 1
    unit       = item.get("unit", "")
    total      = unit_price * quantity
    cv_pct     = nmck_data.get("cv", 0.0) * 100
    name       = item.get("name", "")

    # Заголовок позиции (если их несколько)
    if show_name:
        _add_para(doc, name, bold=True,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  before=120, after=60)

    # ── Таблица контрактов ────────────────────────────────────────
    col_w = [Cm(6.5), Cm(3.5), Cm(7.0)]
    tbl   = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, w in enumerate(col_w):
        tbl.columns[i].width = w

    hdr = tbl.rows[0].cells
    for i, txt in enumerate(("Идентификатор контракта",
                              "Дата заключения контракта",
                              "Цена за единицу товара")):
        _shade_cell(hdr[i], "D9D9D9")
        _cell(hdr[i], txt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for c in used:
        raw_date = str(c.get("Дата заключения контракта", ""))
        d        = raw_date.split(" ")[0].split("T")[0]
        try:
            price_str = _fmt(float(c.get("Цена за единицу", 0)))
        except (ValueError, TypeError):
            price_str = str(c.get("Цена за единицу", ""))
        row = tbl.add_row().cells
        _cell(row[0], str(c.get("Идентификатор контракта", "")))
        _cell(row[1], d,          align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row[2], price_str,  align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_para(doc, before=0, after=60)  # отступ после таблицы

    # ── Заголовок расчёта ────────────────────────────────────────
    _add_para(doc, "Расчет начальной (максимальной) цены контракта:",
              bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=60)

    # ── Таблица расчёта ──────────────────────────────────────────
    rows_data = [
        ("Наименование товара",
         name),
        ("Единица измерения",
         unit or "—"),
        ("Количество",
         str(quantity)),
        ("Коэффициент вариации, %",
         f"{cv_pct:.2f}"),
        ("Средняя арифметическая цена за единицу измерения, руб.",
         _fmt(unit_price)),
        ("Сумма, рассчитанная согласно п. 3.21 МР №567, руб.",
         _fmt(total)),
    ]

    calc_tbl = doc.add_table(rows=len(rows_data), cols=2)
    calc_tbl.style = "Table Grid"
    calc_tbl.columns[0].width = Cm(11.0)
    calc_tbl.columns[1].width = Cm(6.0)

    for i, (label, value) in enumerate(rows_data):
        cells = calc_tbl.rows[i].cells
        _shade_cell(cells[0], "F2F2F2" if i % 2 == 0 else "FFFFFF")
        _shade_cell(cells[1], "F2F2F2" if i % 2 == 0 else "FFFFFF")
        _cell(cells[0], label)
        _cell(cells[1], value, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_para(doc, before=0, after=60)


# ── Параграф об однородности ─────────────────────────────────────

def _cv_paragraph(doc: Document, items: list[dict]) -> None:
    """Параграф о коэффициенте вариации (берётся максимальный из позиций)."""
    cv_values = [
        (item.get("nmck_data") or {}).get("cv", 0.0) * 100
        for item in items
    ]
    max_cv = max(cv_values) if cv_values else 0.0

    if max_cv <= 33:
        text = (
            "Коэффициент вариации цены по закупаемому товару "
            "(по каждому наименованию товара) не превышает 33%, следовательно, "
            "совокупность значений выявленных цен, используемых в расчете "
            "при определении НМЦК, является однородной."
        )
    else:
        text = (
            f"Коэффициент вариации цены составляет {max_cv:.2f}% и превышает 33%, "
            "что свидетельствует о неоднородности совокупности значений цен. "
            "При определении НМЦК применено нормирование в соответствии "
            "с п. 3.20 МР №567."
        )
    _add_para(doc, text, after=120)


# ── Публичный API ─────────────────────────────────────────────────

def build_nmck_docx(items: list[dict], registry_number: str) -> bytes:
    """
    Генерирует DOCX-документ обоснования НМЦК по шаблону МР №567.

    items           — список позиций: name, quantity, unit, unit_price, nmck_data
    registry_number — реестровый номер закупки (генерируется на бэкенде)
    Возвращает байты .docx.
    """
    doc = Document()

    # Поля страницы (ГОСТ: левое 30мм, остальные 20мм)
    sec = doc.sections[0]
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Базовый стиль Normal
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = _SIZE
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Заголовок ─────────────────────────────────────────────────
    _add_para(
        doc,
        "Обоснование начальной (максимальной) цены контракта",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=60,
    )

    # ── Реестровый номер ─────────────────────────────────────────
    p_reg = doc.add_paragraph()
    p_reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_reg, after=240)
    r1 = p_reg.add_run("Реестровый номер закупки: ")
    _set_font(r1, bold=True)
    r2 = p_reg.add_run(registry_number)
    _set_font(r2, bold=True)

    # ── Правовое основание ────────────────────────────────────────
    _add_para(
        doc,
        "В соответствии со статьей 22 Федерального закона от 05 апреля 2013 г. "
        "№ 44-ФЗ «О контрактной системе в сфере закупок товаров, работ, услуг "
        "для обеспечения государственных и муниципальных нужд» начальная "
        "(максимальная) цена контракта (далее - НМЦК) определена и обоснована "
        "посредством применения метода сопоставимых рыночных цен (анализа рынка) "
        "с учетом положений приказа Министерства экономического развития "
        "Российской Федерации от 02.10.2013 № 567 «Об утверждении Методических "
        "рекомендаций по применению методов определения начальной (максимальной) "
        "цены контракта, цены контракта, заключаемого с единственным поставщиком "
        "(подрядчиком, исполнителем)» (далее – МР № 567).",
        after=120,
    )

    _add_para(
        doc,
        "В целях определения НМЦК использовались данные о рыночных ценах "
        "идентичного(ых) и/или сопоставимого(ых) товара(ов), планируемого(ых) "
        "к закупке, полученных путем анализа информации о ценах товаров, работ, "
        "услуг, содержащаяся в контрактах, которые исполнены и по которым не "
        "взыскивались неустойки (штрафы, пени) в связи с неисполнением или "
        "ненадлежащим исполнением обязательств, предусмотренных этими контрактами:",
        after=120,
    )

    # ── Блоки по позициям ────────────────────────────────────────
    multiple = len(items) > 1
    for item in items:
        _item_block(doc, item, show_name=multiple)

    # ── Коэффициент вариации ──────────────────────────────────────
    _cv_paragraph(doc, items)

    # ── Валюта ───────────────────────────────────────────────────
    _add_para(
        doc,
        "Для определения и обоснования начальной (максимальной) цены контракта, "
        "для оплаты поставленного(ых) товара(ов) используется валюта Российской "
        "Федерации – российский рубль.",
        after=240,
    )

    # ── Дата подписания ──────────────────────────────────────────
    _add_para(
        doc,
        "«___»_____________ 20___ г.",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
